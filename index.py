import docx
import os
from docx import Document
import pandas as pd
from docxcompose.composer import Composer

sheet = pd.read_excel("name.xlsx")

def replaceWord(index,name):
    document=Document('temp.docx')
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "i" in run.text:
                run.text=run.text.replace('i', str(index))
            if "name" in run.text:
                run.text=run.text.replace('name',name)
    document.save("ouput/"+name+'.docx')

#for row in sheet.values:
    #replaceWord(row[0],row[1])

original_docx_path='D:/lin/ouput/'
new_docx_path='D:/lin/合并.docx'
def merge_doc(source_file_path_list,target_path_list):
    page_break_doc = Document()
    page_break_doc.add_page_break()
    #定义新文档
    target_doc = Document(source_file_path_list[0])
    target_composer = Composer(target_doc)
    for i in range(len(source_file_path_list)) :
        # 跳过第一个作为模板的文件
        if i==0:
            continue
        # 填充分页符文档
        target_composer.append(page_break_doc)
        # 拼接文档内容
        f = source_file_path_list[i]
        target_composer.append(Document(f))
        # 保存目标文档
    target_composer.save(target_path_list)

source_file_list = os.listdir(original_docx_path)
source_file_list_all = []
for file in source_file_list:
    source_file_list_all.append(original_docx_path+file)
merge_doc(source_file_list_all,new_docx_path)
